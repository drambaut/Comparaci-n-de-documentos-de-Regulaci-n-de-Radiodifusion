import re
from docx import Document

def limpiar_documento_draftable(ruta_entrada, ruta_salida):
    print(f"Iniciando limpieza estructural del documento...")
    try:
        doc = Document(ruta_entrada)
    except Exception as e:
        print(f"Error al abrir {ruta_entrada}: {e}")
        return False

    if not doc.tables:
        print("No se encontraron tablas en el documento.")
        return False
        
    tabla = doc.tables[0]
    total_filas = len(tabla.rows)
    
    patron_relevante = re.compile(
        r'\b(mhz|ghz|khz|hz|mod|add|sup|wrc-23|artículo|art\.|res\.|resolución|colombia)\b', 
        re.IGNORECASE
    )

    filas_a_borrar = []
    for row in tabla.rows[1:]:
        texto_fila = " ".join(cell.text for cell in row.cells)
        if not patron_relevante.search(texto_fila):
            filas_a_borrar.append(row)

    # Eliminación a nivel XML
    for row in filas_a_borrar:
        tr = row._tr
        parent = tr.getparent()
        if parent is not None:
            parent.remove(tr)

    doc.save(ruta_salida)
    
    doc_revisado = Document(ruta_salida)
    filas_finales = len(doc_revisado.tables[0].rows)
    
    print(f"Limpieza completada: De {total_filas} filas brutas a {filas_finales} filas técnicas.")
    return True