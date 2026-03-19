import os
import json
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types

def extraer_datos_limpios(ruta):
    doc = Document(ruta)
    if not doc.tables: return []
    
    tabla = doc.tables[0]
    cambios = []
    for row in tabla.rows[1:]:
        try:
            pag = row.cells[1].text.strip()
            texto = row.cells[6].text.strip()
            notas = row.cells[5].text.strip()
            if texto:
                cambios.append(f"Pág {pag}: {texto} | Contexto: {notas}")
        except IndexError:
            continue
    return cambios

def limpiar_respuesta_json(texto):
    return re.sub(r'```json\n?|```', '', texto).strip()

def analizar_con_experto_ia(lista_cambios, api_key, anio_antiguo, anio_reciente):
    print("Procesando análisis cognitivo con Gemini 2.5 Flash...")
    client = genai.Client(api_key=api_key)
    texto_cambios = "\n".join(lista_cambios)
    
    prompt = f"""
    Actúa como un analista senior de la UIT o la ANE. Redacta un informe ejecutivo técnico basándote en estos cambios detectados entre el Reglamento de Radiocomunicaciones {anio_antiguo} y {anio_reciente}:
    
    {texto_cambios}

    INSTRUCCIONES:
    1. RESUMEN GENERAL: Redacta un párrafo formal explicando los ejes temáticos principales.
    2. SELECCIÓN CRÍTICA: Desprecia cambios de formato. Céntrate en nuevas atribuciones, cambios de límites o notas internacionales.
    3. FORMATO: Prohibido usar asteriscos o markdown. Todo texto plano.

    RESPUESTA REQUERIDA (JSON ESTRICTO):
    {{
        "resumen_ejecutivo": "Tu párrafo aquí...",
        "cambios": [
            {{
                "titulo": "Título breve del cambio",
                "pagina": "número",
                "analisis": "Descripción técnica."
            }}
        ]
    }}
    """
    
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                temperature=0.1,
            ),
        )
        return response.text
    except Exception as e:
        print(f"Error en la IA: {e}")
        return None

def generar_informe_word(json_ia, ruta_salida, anio_antiguo, anio_reciente):
    if not json_ia: return
    try:
        datos = json.loads(limpiar_respuesta_json(json_ia))
    except Exception as e:
        print(f"Error al parsear JSON: {e}")
        return
        
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = p_titulo.add_run(f"Cambios presentados entre el reglamento de radiodifusión del {anio_antiguo} y {anio_reciente}")
    run_titulo.bold = True

    doc.add_paragraph()
    p_resumen = doc.add_paragraph(datos.get("resumen_ejecutivo", ""))
    p_resumen.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()

    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run("Detalle de los cambios realizados:")
    run_sub.bold = True
    doc.add_paragraph()

    for i, cambio in enumerate(datos.get("cambios", []), 1):
        p_item = doc.add_paragraph()
        p_item.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_item = p_item.add_run(f"{i}. {cambio.get('titulo')} (pág. {cambio.get('pagina')})")
        run_item.bold = True
        
        p_desc = doc.add_paragraph(cambio.get('analisis'))
        p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()

    doc.save(ruta_salida)
    print(f"¡Éxito! Informe guardado en: {ruta_salida}")